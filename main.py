import os, io, json, base64, requests
from flask import Flask, render_template_string, request, send_file
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor

app = Flask(__name__)

API_KEY = os.getenv("GOOGLE_API_KEY")

HTML_PAGE = """
<!DOCTYPE html>
<html lang="it">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Convertitore Foto in Word Gratis - OCR Online</title>
    <style>
        body { font-family: system-ui, -apple-system, sans-serif; background: #f4f6f8; margin: 0; padding: 20px; color: #333; }
        .container { max-width: 600px; margin: 0 auto; background: white; padding: 25px; border-radius: 12px; box-shadow: 0 4px 15px rgba(0,0,0,0.1); }
        h1 { font-size: 22px; color: #1a1a1a; text-align: center; }
        p { font-size: 14px; color: #666; line-height: 1.5; text-align: center; }
        input[type="file"] { display: block; margin: 20px auto; }
        button { width: 100%; background: #2563eb; color: white; border: none; padding: 12px; font-size: 16px; border-radius: 6px; cursor: pointer; font-weight: bold; }
        button:hover { background: #1d4ed8; }
        .footer { margin-top: 30px; font-size: 12px; text-align: center; color: #888; }
    </style>
</head>
<body>
    <div class="container">
        <h1>Converti Foto in Word (.docx)</h1>
        <p>Carica un'immagine o scatta una foto al tuo documento.</p>
        <form action="/convert" method="post" enctype="multipart/form-data">
            <input type="file" name="immagine" accept="image/*" required>
            <button type="submit">Converti in Word ora</button>
        </form>
    </div>
    <div class="footer">
        <p>&copy; 2026 Convertitore Foto Word - Strumento OCR Gratuito Online</p>
    </div>
</body>
</html>
"""

@app.route('/')
def home():
    return render_template_string(HTML_PAGE)

@app.route('/test-key')
def test_key():
    if not API_KEY:
        return "Errore: GOOGLE_API_KEY non trovata nelle variabili d'ambiente di Render.", 500
    
    url = f"https://generativelanguage.googleapis.com/v1beta/models?key={API_KEY}"
    res = requests.get(url)
    if res.status_code != 200:
        return f"Errore durante il controllo della chiave API ({res.status_code}): {res.text}", 500
    
    data = res.json()
    models = [m.get("name") for m in data.get("models", []) if "generateContent" in m.get("supportedGenerationMethods", [])]
    return f"<h3>Modelli disponibili per la tua API Key:</h3><pre>{json.dumps(models, indent=2)}</pre>"

@app.route('/convert', methods=['POST'])
def convert():
    file = request.files.get('immagine')
    if not file:
        return "Nessun file caricato", 400
    if not API_KEY:
        return "Errore: GOOGLE_API_KEY non configurata", 500

    try:
        # Trova automaticamente un modello valido tra quelli della tua API Key
        models_url = f"https://generativelanguage.googleapis.com/v1beta/models?key={API_KEY}"
        models_res = requests.get(models_url)
        
        target_model = None
        if models_res.status_code == 200:
            all_models = [m.get("name") for m in models_res.json().get("models", []) if "generateContent" in m.get("supportedGenerationMethods", [])]
            for m in all_models:
                if "flash" in m or "pro" in m:
                    target_model = m
                    break
        
        if not target_model:
            target_model = "models/gemini-1.5-flash"

        # Converte l'immagine
        immagine_pil = Image.open(file.stream)
        buffer = io.BytesIO()
        immagine_pil.convert("RGB").save(buffer, format="JPEG")
        img_b64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

        prompt_text = """
        Analizza l'immagine ed estrai tutto il testo mantenendo la struttura visiva.
        Rispondi ESCLUSIVAMENTE con un JSON valido con questa struttura:
        {
          "blocchi": [
            {
              "tipo": "titolo"|"sottotitolo"|"paragrafo"|"elenco",
              "testo": "testo estratto",
              "grassetto": true|false,
              "corsivo": true|false,
              "dimensione_pt": 11,
              "allineamento": "left"|"center"|"right",
              "colore_hex": "000000",
              "font_suggerito": "Arial"|"Times New Roman"|"Calibri"
            }
          ]
        }
        """

        headers = {'Content-Type': 'application/json'}
        payload = {
            "contents": [{
                "parts": [
                    {"text": prompt_text},
                    {
                        "inline_data": {
                            "mime_type": "image/jpeg",
                            "data": img_b64
                        }
                    }
                ]
            }]
        }

        # Invio richiesta a Google
        url = f"https://generativelanguage.googleapis.com/v1beta/{target_model}:generateContent?key={API_KEY}"
        res = requests.post(url, headers=headers, json=payload)

        if res.status_code != 200:
            return f"Errore Google ({res.status_code}): {res.text}", 500

        res_json = res.json()
        raw_text = res_json['candidates'][0]['content']['parts'][0]['text']
        testo_pulito = raw_text.replace("```json", "").replace("```", "").strip()
        struttura = json.loads(testo_pulito)

        # Genera il file Word
        doc = Document()
        for blocco in struttura.get("blocchi", []):
            tipo = blocco.get("tipo")
            testo = blocco.get("testo", "")
            p = doc.add_paragraph(style='List Bullet') if tipo == "elenco" else doc.add_paragraph()
            run = p.add_run(testo)
            run.font.bold = blocco.get("grassetto", False)
            run.font.italic = blocco.get("corsivo", False)
            run.font.name = blocco.get("font_suggerito", "Calibri")
            run.font.size = Pt(blocco.get("dimensione_pt", 11))
            
            hex_c = blocco.get("colore_hex", "000000").replace("#", "")
            if len(hex_c) == 6:
                try:
                    r, g, b = int(hex_c[0:2], 16), int(hex_c[2:4], 16), int(hex_c[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                except: pass

            align = blocco.get("allineamento", "left")
            p.alignment = 1 if align == "center" else (2 if align == "right" else 0)

        out_buffer = io.BytesIO()
        doc.save(out_buffer)
        out_buffer.seek(0)
        return send_file(out_buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name='documento_convertito.docx')

    except Exception as e:
        return f"Errore durante l'elaborazione: {str(e)}", 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=10000)
    
